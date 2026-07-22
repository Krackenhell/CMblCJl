from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DOC = ROOT / "output" / "doc"
OUT_DOC.mkdir(parents=True, exist_ok=True)

FONT = "Aptos"
DISPLAY = "Aptos Display"

INK = "10211D"
DEEP = "063F35"
GREEN = "0B6B56"
LIME = "B7F400"
BLUE = "63B3FF"
CREAM = "FFFDF4"
MINT = "EEF7F2"
PALE_BLUE = "ECF5FF"
LINE = "D7E3DD"
MUTED = "53635E"
WHITE = "FFFFFF"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run(run, size: float, color: str = INK, bold: bool = False, font: str = FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: str = "8", **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    edge_names = ["top", "left", "bottom", "right", "insideH", "insideV"]
    requested = edges or {name: True for name in edge_names[:4]}
    for name, enabled in requested.items():
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single" if enabled else "nil")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "nil")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_fixed_row_height(row, height_twips: int):
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(height_twips))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Смысл  ·  ")
    set_run(run, 8, MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def configure_page(doc: Document, margins=(1.35, 1.35, 1.2, 1.2)):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    left, right, top, bottom = margins
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.45)
    return section


def configure_styles(doc: Document, normal_size=9.6):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(normal_size)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.06
    for name, size, color in (("Title", 28, DEEP), ("Heading 1", 18, DEEP), ("Heading 2", 12.5, GREEN), ("Heading 3", 10.5, INK)):
        style = styles[name]
        style.font.name = DISPLAY
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DISPLAY)
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(5)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.keep_with_next = True


def add_rich_text(paragraph, text: str, size=9.6, color=INK, font=FONT):
    chunks = re.split(r"(\*\*.*?\*\*)", text)
    for chunk in chunks:
        if not chunk:
            continue
        bold = chunk.startswith("**") and chunk.endswith("**")
        value = chunk[2:-2] if bold else chunk
        run = paragraph.add_run(value)
        set_run(run, size, color, bold, font)


def add_paragraph(doc, text, size=9.6, color=INK, bold=False, align=None, after=4, before=0, keep=False, line=1.06):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_together = keep
    add_rich_text(p, f"**{text}**" if bold else text, size=size, color=color)
    return p


def add_bullet(doc, text, size=9.2, color=INK, level=0, after=2):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.45 + level * 0.4)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.02
    add_rich_text(p, text, size=size, color=color)
    return p


def add_callout(doc, title, body, fill=MINT, accent=GREEN, title_size=10.5, body_size=9.4):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(0.18)
    table.columns[1].width = Cm(17.7)
    table.cell(0, 0).width = Cm(0.18)
    table.cell(0, 1).width = Cm(17.7)
    remove_table_borders(table)
    shade_cell(table.cell(0, 0), accent)
    shade_cell(table.cell(0, 1), fill)
    for cell in table.row_cells(0):
        set_cell_margins(cell, top=120, bottom=120, start=150, end=150)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = table.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run(r, title_size, DEEP, True, DISPLAY)
    p2 = table.cell(0, 1).add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.04
    add_rich_text(p2, body, size=body_size, color=INK)
    return table


def add_section_band(doc, index, title, subtitle=None):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(2.0)
    table.columns[1].width = Cm(16.2)
    table.cell(0, 0).width = Cm(2.0)
    table.cell(0, 1).width = Cm(16.2)
    remove_table_borders(table)
    shade_cell(table.cell(0, 0), LIME)
    shade_cell(table.cell(0, 1), DEEP)
    for c in table.row_cells(0):
        set_cell_margins(c, top=140, bottom=140, start=180, end=180)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = table.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(index)
    set_run(r, 14, DEEP, True, DISPLAY)
    p2 = table.cell(0, 1).paragraphs[0]
    r2 = p2.add_run(title)
    set_run(r2, 17, WHITE, True, DISPLAY)
    if subtitle:
        r3 = p2.add_run("  ·  " + subtitle)
        set_run(r3, 9, "B9D6CD", False, FONT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def base_properties(doc: Document, title: str, subject: str):
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "Иван Евгеньевич Сороковых"
    props.keywords = "Junior ML Contest, AI Talent Hub, Смысл, ИИ в образовании"


def make_motivation_letter() -> Path:
    doc = Document()
    section = configure_page(doc, margins=(1.65, 1.65, 1.25, 1.15))
    configure_styles(doc, normal_size=9.7)
    base_properties(doc, "Мотивационное письмо — Иван Сороковых", "Поступление на программу «Искусственный интеллект» AI Talent Hub")

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("AI TALENT HUB  ·  ИСКУССТВЕННЫЙ ИНТЕЛЛЕКТ")
    set_run(r, 7.7, GREEN, True)
    add_page_field(section.footer.paragraphs[0])

    band = doc.add_table(rows=1, cols=2)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.autofit = False
    band.columns[0].width = Cm(12.8)
    band.columns[1].width = Cm(4.6)
    band.cell(0, 0).width = Cm(12.8)
    band.cell(0, 1).width = Cm(4.6)
    remove_table_borders(band)
    shade_cell(band.cell(0, 0), DEEP)
    shade_cell(band.cell(0, 1), LIME)
    for cell in band.row_cells(0):
        set_cell_margins(cell, top=160, bottom=160, start=200, end=200)
    p = band.cell(0, 0).paragraphs[0]
    r = p.add_run("МОТИВАЦИОННОЕ ПИСЬМО")
    set_run(r, 16, WHITE, True, DISPLAY)
    p = band.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026")
    set_run(r, 14, DEEP, True, DISPLAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("Иван Евгеньевич Сороковых")
    set_run(r, 21, DEEP, True, DISPLAY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Кандидат на программу «Искусственный интеллект» · Москва")
    set_run(r, 9.5, MUTED)

    source = (ROOT / "docs" / "MOTIVATION_LETTER_AI_TALENT_HUB.md").read_text(encoding="utf-8")
    blocks = [block.strip() for block in re.split(r"\n\s*\n", source) if block.strip()]
    body_started = False
    for block in blocks:
        if block.startswith("Уважаемые члены"):
            body_started = True
        if not body_started:
            continue
        if block.startswith("**Иван Сороковых**"):
            break
        text = block.replace("  \n", " ").replace("\n", " ")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0.55)
        p.paragraph_format.space_after = Pt(4.3)
        p.paragraph_format.line_spacing = 1.03
        p.paragraph_format.widow_control = True
        add_rich_text(p, text, size=9.7)

    sign = doc.add_table(rows=1, cols=2)
    sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign.autofit = False
    sign.columns[0].width = Cm(10)
    sign.columns[1].width = Cm(7.2)
    sign.cell(0, 0).width = Cm(10)
    sign.cell(0, 1).width = Cm(7.2)
    remove_table_borders(sign)
    p = sign.cell(0, 0).paragraphs[0]
    r = p.add_run("С уважением,")
    set_run(r, 9, MUTED)
    p2 = sign.cell(0, 0).add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run("Иван Сороковых")
    set_run(r, 11.5, DEEP, True, DISPLAY)
    p = sign.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("sorokovyh04@mail.ru")
    set_run(r, 9, GREEN)

    out = OUT_DOC / "Motivation_Letter_Ivan_Sorokovykh_AI_Talent_Hub.docx"
    doc.save(out)
    return out


def make_project_brief() -> Path:
    doc = Document()
    section = configure_page(doc, margins=(1.35, 1.35, 1.15, 1.05))
    configure_styles(doc, normal_size=9.15)
    base_properties(doc, "VivaTrace — краткое описание проекта", "Junior ML Contest 2026 · AI в образовании")
    add_page_field(section.footer.paragraphs[0])

    # PAGE 1 — executive cover
    top = doc.add_table(rows=1, cols=2)
    top.alignment = WD_TABLE_ALIGNMENT.CENTER
    top.autofit = False
    top.columns[0].width = Cm(12.7)
    top.columns[1].width = Cm(5.3)
    top.cell(0, 0).width = Cm(12.7)
    top.cell(0, 1).width = Cm(5.3)
    remove_table_borders(top)
    shade_cell(top.cell(0, 0), DEEP)
    shade_cell(top.cell(0, 1), LIME)
    for c in top.row_cells(0):
        set_cell_margins(c, top=125, bottom=125, start=180, end=180)
    p = top.cell(0, 0).paragraphs[0]
    r = p.add_run("JUNIOR ML CONTEST 2026")
    set_run(r, 9, WHITE, True)
    p = top.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI В ОБРАЗОВАНИИ")
    set_run(r, 8.5, DEEP, True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("Viva")
    set_run(r, 34, DEEP, True, DISPLAY)
    r = p.add_run("Trace")
    set_run(r, 34, GREEN, True, DISPLAY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("От ответа студента — к доказательству понимания\nи плану следующего занятия")
    set_run(r, 17, INK, True, DISPLAY)

    add_callout(
        doc,
        "Проблема",
        "Одинаковая программа не означает одинаковое понимание. Тест показывает результат, но не всегда показывает знание; преподаватель узнаёт об общих пробелах слишком поздно.",
        fill="F6F4EA",
        accent=LIME,
        title_size=11,
        body_size=10,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    add_paragraph(doc, "Замкнутый учебный цикл", size=12.5, color=DEEP, bold=True, before=2, after=4)
    flow = doc.add_table(rows=1, cols=6)
    flow.alignment = WD_TABLE_ALIGNMENT.CENTER
    flow.autofit = False
    labels = ["1\nЗадание", "2\nMicro-Viva", "3\nПерсональный\nмаршрут", "4\nПрактическая\nмиссия", "5\nПульс\nгруппы", "6\nПлан\nпреподавателю"]
    fills = [DEEP, GREEN, "168468", BLUE, "317FC6", DEEP]
    for index, label in enumerate(labels):
        cell = flow.cell(0, index)
        cell.width = Cm(3)
        shade_cell(cell, fills[index])
        set_cell_margins(cell, top=120, bottom=120, start=70, end=70)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(cell, color=WHITE, size="14")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 0.95
        r = p.add_run(label)
        set_run(r, 8.2, WHITE, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Код проверяет то, что можно проверить точно.  ·  Локальная LLM объясняет, задаёт вопросы и персонализирует.")
    set_run(r, 9.2, GREEN, True)

    add_paragraph(doc, "Измеримый MVP", size=12.5, color=DEEP, bold=True, after=4)
    metrics = [
        ("90", "заданий B2\n10 тем × 9 вариантов", MINT),
        ("100", "автоматических тестов\nосновного учебного цикла", PALE_BLUE),
        ("14 / 14", "ответов Viva\nв экспертном диапазоне", "F4F8DF"),
        ("LOCAL-FIRST", "Qwen 2.5 · основной цикл\nбез облачного API", "F0F3F2"),
    ]
    grid = doc.add_table(rows=2, cols=2)
    grid.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid.autofit = False
    for idx, (value, label, fill) in enumerate(metrics):
        row, col = divmod(idx, 2)
        cell = grid.cell(row, col)
        cell.width = Cm(9)
        shade_cell(cell, fill)
        set_cell_border(cell, color=WHITE, size="20")
        set_cell_margins(cell, top=135, bottom=135, start=180, end=180)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(value)
        set_run(r, 17 if value != "LOCAL-FIRST" else 13, DEEP, True, DISPLAY)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 0.96
        r = p2.add_run(label)
        set_run(r, 8.8, MUTED)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    add_callout(
        doc,
        "Не ещё один чат с LLM",
        "VivaTrace разделяет ответственность: предметный модуль проверяет закрытые ответы, LLM работает с семантикой, а освоение собирается из нескольких наблюдений. Персонализация AI сохраняется без отказа от воспроизводимости обычного ПО.",
        fill=DEEP,
        accent=LIME,
        title_size=11,
        body_size=9.2,
    )
    # recolor text in the dark callout
    dark_cell = doc.tables[-1].cell(0, 1)
    for paragraph in dark_cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = rgb(WHITE if run.bold else "D7ECE5")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ИВАН СОРОКОВЫХ")
    set_run(r, 9, DEEP, True)
    r = p.add_run("  ·  Product vision  ·  AI-assisted development  ·  тестирование и валидация")
    set_run(r, 8.7, MUTED)

    doc.add_page_break()

    # PAGE 2 — problem, data, journey, architecture
    add_section_band(doc, "01", "ЗАДАЧА, ДАННЫЕ И ПУТЬ К АРХИТЕКТУРЕ", "от прототипа к проверяемому циклу")
    add_paragraph(doc, "Продуктовая гипотеза", size=12.3, color=DEEP, bold=True, after=2)
    add_paragraph(
        doc,
        "Первичный пользователь — студент в неоднородной группе; второй — преподаватель, которому нужно решить, продолжать программу или закрепить общий пробел. Гипотеза: короткая проверка объяснения сразу после задания лучше отделяет воспроизведённый ответ от понимания, а агрегированные ошибки позволяют точнее планировать следующую пару.",
        size=9.2,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        after=3,
    )
    add_callout(
        doc,
        "Разрыв существующих решений",
        "LMS-тест проверяет артефакт ответа, а универсальный LLM-чат не хранит учебный контур группы и может уверенно галлюцинировать. VivaTrace соединяет индивидуальную диагностику и решение преподавателя в одном цикле.",
        fill="F6F4EA",
        accent=LIME,
        title_size=9.8,
        body_size=8.8,
    )

    add_paragraph(doc, "Как развивался проект", size=12.3, color=DEEP, bold=True, before=4, after=3)
    journey = [
        ("01", "Статичный прототип", "Проверили ценность связки «задание — Viva — пульс группы», но отказались считать демо-цифры доказательством."),
        ("02", "Реальные профили и события", "Попытки разных студентов стали сохраняться в SQLite и сразу менять аналитику преподавателя."),
        ("03", "Local-first", "Основной цикл перенесён на Qwen 2.5 GGUF: проверяющему не нужны ключи, учебные данные остаются локально."),
        ("04", "LLM перестала быть предметным судьёй", "После ложных 70–85% на случайных ответах закрытые задания переданы deterministic grader; модель получает проверенные факты."),
        ("05", "Калибровка семантики", "Добавлены экспертные диапазоны, no-knowledge/tautology checks, benchmark и сравнение алгоритмов освоения."),
        ("06", "Перенос навыка и контур группы", "Практические миссии, голос, прогресс и преподавательские рекомендации строятся на подтверждённых компонентах ошибок."),
    ]
    jt = doc.add_table(rows=len(journey), cols=2)
    jt.alignment = WD_TABLE_ALIGNMENT.CENTER
    jt.autofit = False
    jt.columns[0].width = Cm(1.15)
    jt.columns[1].width = Cm(16.8)
    remove_table_borders(jt)
    for idx, (num, title, body) in enumerate(journey):
        c0, c1 = jt.cell(idx, 0), jt.cell(idx, 1)
        c0.width = Cm(1.15)
        c1.width = Cm(16.8)
        shade_cell(c0, LIME if idx in (3, 4) else MINT)
        shade_cell(c1, "FAFCFB" if idx % 2 == 0 else WHITE)
        set_cell_margins(c0, top=70, bottom=70, start=60, end=60)
        set_cell_margins(c1, top=70, bottom=70, start=120, end=120)
        c0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c0.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(num)
        set_run(r, 8.5, DEEP, True)
        p = c1.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(title + ". ")
        set_run(r, 8.75, DEEP, True)
        r = p.add_run(body)
        set_run(r, 8.45, INK)

    add_paragraph(doc, "Данные", size=12.3, color=DEEP, bold=True, before=4, after=2)
    data_table = doc.add_table(rows=2, cols=2)
    data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    data_table.autofit = False
    data_items = [
        ("Учебный банк", "90 заданий B2 по 10 темам; 71 структурированное задание поддерживает точную поэлементную проверку."),
        ("Диагностика", "14 экспертно размеченных сценариев Viva: парафразы, частичные, случайные и терминологически неточные ответы."),
        ("Практика", "10 сценариев переноса навыка и 31 контрольный пример; live-попытки связаны со студентом, темой и trace ID."),
        ("Knowledge tracing", "180 синтетических студентов и 660 test events — только proof-of-pipeline, не доказательство образовательного эффекта."),
    ]
    for idx, (title, body) in enumerate(data_items):
        row, col = divmod(idx, 2)
        cell = data_table.cell(row, col)
        cell.width = Cm(9)
        shade_cell(cell, MINT if idx % 2 == 0 else PALE_BLUE)
        set_cell_border(cell, color=WHITE, size="18")
        set_cell_margins(cell, top=100, bottom=100, start=130, end=130)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(title)
        set_run(r, 8.8, DEEP, True)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 0.98
        r = p2.add_run(body)
        set_run(r, 8.1, INK)

    add_paragraph(doc, "Архитектурный инвариант", size=12.3, color=DEEP, bold=True, before=4, after=2)
    arch = doc.add_table(rows=1, cols=5)
    arch.alignment = WD_TABLE_ALIGNMENT.CENTER
    arch.autofit = False
    arch_items = [
        ("1", "Предметный\ngrader"),
        ("2", "Проверенные\nошибки"),
        ("3", "Qwen:\nViva + разбор"),
        ("4", "Evidence\nfusion"),
        ("5", "Маршрут +\nпульс группы"),
    ]
    for idx, (num, label) in enumerate(arch_items):
        cell = arch.cell(0, idx)
        shade_cell(cell, DEEP if idx in (0, 2, 4) else GREEN)
        set_cell_border(cell, color=WHITE, size="18")
        set_cell_margins(cell, top=100, bottom=100, start=60, end=60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 0.94
        r = p.add_run(num + "\n" + label)
        set_run(r, 8, WHITE, True)
    add_paragraph(
        doc,
        "Модель, hash запроса, этап, длительность и результат сохраняются в журнале. Невалидный JSON не может молча пройти дальше; голосовой режим изолирован и не ломает основной учебный контур.",
        size=8.35,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=3,
        after=0,
    )

    doc.add_page_break()

    # PAGE 3 — experiments and criteria
    add_section_band(doc, "02", "МЕТОДЫ, РЕЗУЛЬТАТЫ И ЦЕННОСТЬ", "измеримость вместо демо-эффекта")
    add_paragraph(doc, "Что проверено", size=12.3, color=DEEP, bold=True, after=3)
    rows = [
        ("Автоматические тесты", "100 passed", "Регрессии grader, Viva, кабинетов, миссий и голоса"),
        ("Semantic Viva", "14/14", "В экспертных диапазонах; 0 false accept и 0 false reject"),
        ("Practical Mission", "20/20 + 11/11", "Deterministic и live Qwen; неподтверждённые ошибки отбрасываются"),
        ("Synthetic proof-of-pipeline", "0,751 → 0,906", "Balanced accuracy assignment-only → assignment+Viva; не эффект обучения"),
        ("Knowledge tracing", "EMA Brier 0,2208", "BKT: 0,2298; EMA оставлен более калиброванным baseline"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(4.1), Cm(3.2), Cm(10.6)]
    headers = ["Эксперимент", "Результат", "Корректная интерпретация"]
    for idx, value in enumerate(headers):
        cell = table.cell(0, idx)
        cell.width = widths[idx]
        shade_cell(cell, DEEP)
        set_cell_margins(cell, top=100, bottom=100, start=120, end=120)
        p = cell.paragraphs[0]
        r = p.add_run(value)
        set_run(r, 8.2, WHITE, True)
    set_repeat_table_header(table.rows[0])
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].width = widths[idx]
            shade_cell(cells[idx], MINT if row_idx % 2 == 0 else WHITE)
            set_cell_border(cells[idx], color=LINE, size="6")
            set_cell_margins(cells[idx], top=80, bottom=80, start=110, end=110)
            p = cells[idx].paragraphs[0]
            p.paragraph_format.line_spacing = 0.98
            r = p.add_run(value)
            set_run(r, 8.2, DEEP if idx == 1 else INK, idx == 1)

    add_paragraph(
        doc,
        "Честное ограничение: средняя задержка локальной семантической оценки на текущем CPU — 16,8 с. Следующая инженерная задача — очередь инференса, кеширование и меньшая специализированная модель.",
        size=8.4,
        color=MUTED,
        before=3,
        after=3,
    )

    add_paragraph(doc, "Соответствие критериям конкурса", size=12.3, color=DEEP, bold=True, after=3)
    criteria = [
        ("Разработка и инженерия", "Модульная архитектура, SQLite, Docker/Compose, CI, local inference, trace-журнал, отказоустойчивый JSON-контур и 100 тестов."),
        ("Data Science", "Экспертные диапазоны, false accept/reject, Brier, log loss, ECE и baseline-модели. Учебные, диагностические и синтетические данные разделены."),
        ("Применение AI", "LLM отвечает за семантику, Viva и персонализацию, но ограничена проверенными фактами. AI-агенты использованы в исследовании, коде, тестах и UX."),
        ("Продуктовое мышление", "Студент получает разбор и новую практику; преподаватель — реальные пробелы группы и тему занятия. XP не подменяет освоение."),
    ]
    ct = doc.add_table(rows=2, cols=2)
    ct.alignment = WD_TABLE_ALIGNMENT.CENTER
    ct.autofit = False
    fills = [MINT, PALE_BLUE, "F4F8DF", "F6F4EA"]
    for idx, (title, body) in enumerate(criteria):
        row, col = divmod(idx, 2)
        cell = ct.cell(row, col)
        cell.width = Cm(9)
        shade_cell(cell, fills[idx])
        set_cell_border(cell, color=WHITE, size="18")
        set_cell_margins(cell, top=110, bottom=110, start=140, end=140)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(title)
        set_run(r, 9, DEEP, True)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 0.98
        r = p2.add_run(body)
        set_run(r, 8.15, INK)

    add_callout(
        doc,
        "Личный вклад автора",
        "Иван Сороковых сформулировал проблему и продуктовую гипотезу, спроектировал пользовательский цикл, задавал AI-агенту требования и критерии приёмки, тестировал разные профили, находил системные ошибки оценивания и принимал решения о переходе к гибридной архитектуре. Код создавался в AI-assisted режиме — это открыто отражено в Git и документации.",
        fill="F4F8DF",
        accent=LIME,
        title_size=9.8,
        body_size=8.35,
    )

    add_paragraph(doc, "Следующий доказуемый шаг", size=11.4, color=DEEP, bold=True, before=4, after=1)
    add_paragraph(
        doc,
        "Пилот на 10–20 студентах: экспертная разметка, межэкспертное согласие, latency/error budget, сравнение assignment-only и assignment+Viva, интервью преподавателей. Затем — RAG по материалам конкретного курса и LMS-интеграция.",
        size=8.6,
        after=3,
    )
    add_callout(
        doc,
        "Почему VivaTrace достоин победы",
        "Это не декоративная оболочка над LLM, а работающий и аудируемый цикл «доказательство понимания → действие студента → решение преподавателя». Проект прошёл путь от эффектной демо-идеи к системе, которая знает, где модели нельзя доверять, измеряет качество и честно показывает границы результата.",
        fill=DEEP,
        accent=LIME,
        title_size=10.5,
        body_size=8.8,
    )
    dark_cell = doc.tables[-1].cell(0, 1)
    for paragraph in dark_cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = rgb(WHITE if run.bold else "D7ECE5")

    out = OUT_DOC / "VivaTrace_Project_Brief_JMLC_2026.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    paths = [make_motivation_letter(), make_project_brief()]
    for path in paths:
        print(path)
